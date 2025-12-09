from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, session
from flask import send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import json
import os

app = Flask(__name__)
app.config.from_object('config.DevelopmentConfig')

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(120), nullable=False)
    user_type = db.Column(db.String(20), nullable=False)  #student, admin, superadmin
    name = db.Column(db.String(100))
    email = db.Column(db.String(120))
    sig = db.Column(db.String(50))
    position = db.Column(db.String(50))
    
    student_number = db.Column(db.String(20))
    year = db.Column(db.Integer)
    section = db.Column(db.String(5))
    major = db.Column(db.String(20))
    status = db.Column(db.String(20), default='pending')  #pindeng, aprobd, rijictid

class Event(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text)
    date = db.Column(db.Date, nullable=False)
    time = db.Column(db.Time)
    room = db.Column(db.String(20))
    sig = db.Column(db.String(50), nullable=False)
    event_type = db.Column(db.String(20))  # eveny o meeting
    status = db.Column(db.String(20), default='pending')  #pending, approved, rejected
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'))
    feedback = db.Column(db.Text)

class Officer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    position = db.Column(db.String(50), nullable=False)
    sig = db.Column(db.String(50), nullable=False)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))




@app.route('/')
def index():
    if current_user.is_authenticated:
        if current_user.user_type == 'student':
            return redirect(url_for('student_dashboard'))
        elif current_user.user_type == 'admin':
            return redirect(url_for('admin_dashboard'))
        elif current_user.user_type == 'superadmin':
            return redirect(url_for('superadmin_dashboard'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user_type = request.form.get('userType')
        
        user = User.query.filter_by(username=username, user_type=user_type).first()
        
        if user and check_password_hash(user.password, password):
            login_user(user)
            flash('Logged in successfully!', 'success')
            
            if user_type == 'student':
                return redirect(url_for('student_dashboard'))
            elif user_type == 'admin':
                return redirect(url_for('admin_dashboard'))
            elif user_type == 'superadmin':
                return redirect(url_for('superadmin_dashboard'))
        else:
            flash('Invalid credentials!', 'error')
    
    return render_template('login.html')

@app.route('/signup', methods=['POST'])
def signup():
    if request.method == 'POST':
        data = request.get_json()
        
        if User.query.filter_by(student_number=data['studentNumber']).first():
            return jsonify({'success': False, 'message': 'Student number already exists!'})
        
        if User.query.filter_by(username=data['studentNumber']).first():
            return jsonify({'success': False, 'message': 'Username already exists!'})
        
        user = User(
            username=data['studentNumber'],
            password=generate_password_hash(data['password']),
            user_type='student',
            name=data['name'],
            email=data['email'],
            student_number=data['studentNumber'],
            year=int(data['year']),
            section=data['section'],
            major=data.get('major', ''),
            sig=data['sig'],
            status='pending'
        )
        
        db.session.add(user)
        db.session.commit()
        
        return jsonify({'success': True, 'message': 'Application submitted! Please wait for Officer approval.'})

@app.route('/admin/applications')
@login_required
def admin_applications():
    if current_user.user_type != 'admin':
        return redirect(url_for('login'))
    
    #no no don't touch me there, this is my no no square
    print(f"\ndebugbug a berna")
    print(f"Admin: {current_user.name} ({current_user.sig})")
    
    #pending applications
    pending_students = User.query.filter_by(
        sig=current_user.sig, 
        user_type='student', 
        status='pending'
    ).order_by(User.id.desc()).all()
    
    #rejected applications
    rejected_students = User.query.filter_by(
        sig=current_user.sig, 
        user_type='student', 
        status='rejected'
    ).order_by(User.id.desc()).all()
    
    #aLL students
    all_students = User.query.filter_by(
        sig=current_user.sig, 
        user_type='student'
    ).all()
    
    print(f"total students in SIG: {len(all_students)}")
    for student in all_students:
        print(f"  - {student.name}: status={student.status}")
    
    print(f"pending students count: {len(pending_students)}")
    print(f"rejected students count: {len(rejected_students)}")
    print("The End\n")
    #wag na lang galawin ito ^ kase pang debug yan
    
    return render_template('admin_applications.html', 
                         pending_students=pending_students,
                         rejected_students=rejected_students)

@app.route('/admin/delete_application', methods=['POST'])
@login_required
def delete_application():
    if current_user.user_type != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized!'})
    
    data = request.get_json()
    student_id = data['student_id']
    
    student = User.query.get(student_id)
    
    if not student or student.sig != current_user.sig:
        return jsonify({'success': False, 'message': 'Student not found or wrong SIG!'})
    
    if student.status != 'rejected':
        return jsonify({'success': False, 'message': 'Only rejected applications can be deleted!'})
    
    #pang display lang sa alert
    student_name = student.name
    student_number = student.student_number
    
    db.session.delete(student)
    db.session.commit()
    
    return jsonify({'success': True, 'message': f'Application for {student_name} ({student_number}) deleted successfully!'})

@app.route('/admin/handle_application', methods=['POST'])
@login_required
def handle_application():
    if current_user.user_type != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized!'})
    
    data = request.get_json()
    student_id = data['student_id']
    action = data['action']  #approve, reject, o reconsider
    
    
    student = User.query.get(student_id)
    
    if not student:
        print(f"ERROR: Student not found with ID {student_id}")
        return jsonify({'success': False, 'message': 'Student not found!'})
    
    
    if student.sig != current_user.sig:
        return jsonify({'success': False, 'message': 'Wrong SIG!'})
    
    if action == 'reconsider':
        student.status = 'pending'
        message = 'Application moved back to pending!'
        print(f"changed status to: pending")
    else:
        new_status = action + 'd'  #approved o rejected
        student.status = new_status
        message = f'Application {action}d successfully!'
        print(f"changed status to: {new_status}")
    
    try:
        db.session.commit()
        print(f"SUCCESS: Database committed")
        
        #refresh at verify pagkasave
        db.session.refresh(student)
        print(f"verified status: {student.status}")
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'message': f'Database error: {str(e)}'})
    
    return jsonify({'success': True, 'message': message})

@app.route('/superadmin/applications')
@login_required
def superadmin_applications():
    if current_user.user_type != 'superadmin':
        return redirect(url_for('login'))
    
    #lahat ng pending application
    pending_students = User.query.filter_by(user_type='student', status='pending').all()
    
    return render_template('superadmin_applications.html', pending_students=pending_students)

@app.route('/student/dashboard')
@login_required
def student_dashboard():
    if current_user.user_type != 'student':
        return redirect(url_for('login'))

    if current_user.status != 'approved':
        return render_template('student_pending.html', status=current_user.status or 'pending')
    
    events = Event.query.filter_by(sig=current_user.sig, status='approved').all()
    
    officers = Officer.query.filter_by(sig=current_user.sig).all()
    
    members_count = User.query.filter_by(sig=current_user.sig, user_type='student', status='approved').count()
    
    return render_template('student_dashboard.html', 
                         events=events, 
                         officers=officers,
                         members_count=members_count)

@app.route('/admin/dashboard')
@login_required
def admin_dashboard():
    if current_user.user_type != 'admin':
        return redirect(url_for('login'))
    
    approved_members_count = User.query.filter_by(
        sig=current_user.sig, 
        user_type='student', 
        status='approved'
    ).count()
    
    pending_count = User.query.filter_by(
        sig=current_user.sig, 
        user_type='student', 
        status='pending'
    ).count()
    
    officers = Officer.query.filter_by(sig=current_user.sig).all()
    
    events = Event.query.filter_by(sig=current_user.sig, status='approved').all()
    
    return render_template('admin_dashboard.html',
                         approved_members_count=approved_members_count,
                         pending_count=pending_count,
                         officers=officers,
                         events=events)

@app.route('/superadmin/dashboard')
@login_required
def superadmin_dashboard():
    if current_user.user_type != 'superadmin':
        return redirect(url_for('login'))
    
    sigs = ['CodEx', 'Netac', 'Source Code', 'Robotix', 'Graphicos']
    stats = {}
    total_students = 0
    
    for sig in sigs:
        count = User.query.filter_by(sig=sig, user_type='student').count()
        stats[sig] = count
        total_students += count
    
    pending_requests = Event.query.filter_by(status='pending').all()
    
    return render_template('superadmin_dashboard.html',
                         stats=stats,
                         total_students=total_students,
                         pending_requests=pending_requests)

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    if request.method == 'POST':
        current_user.name = request.form.get('name')
        current_user.email = request.form.get('email')
        
        if current_user.user_type == 'student':
            current_user.year = int(request.form.get('year'))
            current_user.section = request.form.get('section')
            current_user.major = request.form.get('major')
            current_user.sig = request.form.get('sig')
        
        db.session.commit()
        flash('Profile updated successfully!', 'success')
        return redirect(url_for('profile'))
    
    return render_template('profile.html')

@app.route('/change_password', methods=['POST'])
@login_required
def change_password():
    current_password = request.form.get('currentPassword')
    new_password = request.form.get('newPassword')
    
    if check_password_hash(current_user.password, current_password):
        current_user.password = generate_password_hash(new_password)
        db.session.commit()
        flash('Password changed successfully!', 'success')
    else:
        flash('Current password is incorrect!', 'error')
    
    return redirect(url_for('profile'))

@app.route('/members')
@login_required
def members():
    year_filter = request.args.get('year', '')
    section_filter = request.args.get('section', '')
    major_filter = request.args.get('major', '')
    sig_filter = request.args.get('sig', '')
    search_query = request.args.get('search', '')
    
    if current_user.user_type == 'student':
        query = User.query.filter_by(
            sig=current_user.sig, 
            user_type='student', 
            status='approved'
        )
    elif current_user.user_type == 'admin':
        query = User.query.filter_by(
            sig=current_user.sig, 
            user_type='student'
        ).filter(
            User.status.in_(['approved', 'pending'])
        )
    else:  #superadmin
        query = User.query.filter_by(user_type='student')
    
    #filter
    if year_filter:
        query = query.filter_by(year=int(year_filter))
    
    if section_filter:
        query = query.filter_by(section=section_filter)
    
    if major_filter:
        query = query.filter_by(major=major_filter)
    
    if sig_filter and current_user.user_type == 'superadmin':
        query = query.filter_by(sig=sig_filter)
    
    #pang search
    if search_query:
        query = query.filter(User.student_number.contains(search_query))
    
    #execute query
    members = query.order_by(User.name).all()
    
    #para sa dropdown
    unique_years = sorted({m.year for m in User.query.filter_by(user_type='student').all() if m.year})
    unique_sections = sorted({m.section for m in User.query.filter_by(user_type='student').all() if m.section})
    unique_majors = sorted({m.major for m in User.query.filter_by(user_type='student').all() if m.major and m.major != ''})
    unique_sigs = sorted({m.sig for m in User.query.filter_by(user_type='student').all() if m.sig})
    
    return render_template('members.html', 
                         members=members,
                         year_filter=year_filter,
                         section_filter=section_filter,
                         major_filter=major_filter,
                         sig_filter=sig_filter,
                         search_query=search_query,
                         unique_years=unique_years,
                         unique_sections=unique_sections,
                         unique_majors=unique_majors,
                         unique_sigs=unique_sigs)

@app.route('/events')
@login_required
def events():
    if current_user.user_type == 'student':
        events = Event.query.filter_by(sig=current_user.sig, status='approved').all()
    elif current_user.user_type == 'admin':
        events = Event.query.filter_by(sig=current_user.sig).all()
    else:  #superadmin
        events = Event.query.all()
    
    return render_template('events.html', events=events)

@app.route('/submit_request', methods=['POST'])
@login_required
def submit_request():
    if current_user.user_type != 'admin':
        return jsonify({'success': False, 'message': 'Unauthorized!'})
    
    data = request.get_json()
    
    #checker ng date conflict
    if data['type'] == 'event':
        existing_event = Event.query.filter_by(date=datetime.strptime(data['date'], '%Y-%m-%d').date(), status='approved').first()
        if existing_event:
            return jsonify({'success': False, 'message': 'Date conflict with existing event!'})
    
    #checker ng time and room conflict
    elif data['type'] == 'meeting':
        existing_meeting = Event.query.filter_by(
            date=datetime.strptime(data['date'], '%Y-%m-%d').date(),
            time=datetime.strptime(data['time'], '%H:%M').time(),
            room=data['room'],
            status='approved'
        ).first()
        if existing_meeting:
            return jsonify({'success': False, 'message': 'Time and room conflict with existing meeting!'})
    
    event = Event(
        title=data['title'],
        description=data['description'],
        date=datetime.strptime(data['date'], '%Y-%m-%d').date(),
        time=datetime.strptime(data['time'], '%H:%M').time() if data.get('time') else None,
        room=data.get('room'),
        sig=current_user.sig,
        event_type=data['type'],
        created_by=current_user.id
    )
    
    db.session.add(event)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Request submitted successfully!'})

@app.route('/handle_request', methods=['POST'])
@login_required
def handle_request():
    if current_user.user_type != 'superadmin':
        return jsonify({'success': False, 'message': 'Unauthorized!'})
    
    data = request.get_json()
    event = Event.query.get(data['event_id'])
    
    if event:
        event.status = data['action']
        event.feedback = data.get('feedback', '')
        db.session.commit()
        return jsonify({'success': True, 'message': f'Request {data["action"]}!'})
    
    return jsonify({'success': False, 'message': 'Event not found!'})

@app.route('/officers', methods=['GET', 'POST'])
@login_required
def officers():
    if current_user.user_type != 'admin':
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        data = request.get_json()
        
        Officer.query.filter_by(sig=current_user.sig).delete()
        
        for officer_data in data['officers']:
            officer = Officer(
                name=officer_data['name'],
                position=officer_data['position'],
                sig=current_user.sig
            )
            db.session.add(officer)
        
        db.session.commit()
        return jsonify({'success': True, 'message': 'Officers updated successfully!'})
    
    officers = Officer.query.filter_by(sig=current_user.sig).all()
    return render_template('officers.html', officers=officers)

@app.route('/requests')
@login_required
def requests():
    if current_user.user_type == 'admin':
        pending_requests = Event.query.filter_by(sig=current_user.sig, status='pending').all()
        return render_template('requests.html', pending_requests=pending_requests)
    elif current_user.user_type == 'superadmin':
        return redirect(url_for('superadmin_dashboard'))
    else:
        return redirect(url_for('login'))

@app.route('/reports')
@login_required
def reports():
    if current_user.user_type in ['admin', 'superadmin']:
        return render_template('reports.html')
    else:
        return redirect(url_for('login'))

def create_report_document():
    """Create a document using the template"""
    try:
        
        template_path = 'templates_docs/template.docx'
        doc = Document(template_path)
        
        if len(doc.paragraphs) > 0:
            if 'SIGMS REPORT' in doc.paragraphs[0].text:
                pass
    except Exception as e:
        print(f"Template not found or error: {e}")
        doc = Document()
    
    return doc

#download rouites
@app.route('/download/member_list')
@login_required
def download_member_list():
    if current_user.user_type != 'admin':
        return redirect(url_for('login'))
    
    #get sig mems
    members = User.query.filter_by(
        sig=current_user.sig, 
        user_type='student', 
        status='approved'
    ).order_by(User.id).all()
    
    year_filter = request.args.get('year', '')
    section_filter = request.args.get('section', '')
    major_filter = request.args.get('major', '')
    search_query = request.args.get('search', '')
    
    query = User.query.filter_by(
        sig=current_user.sig, 
        user_type='student', 
        status='approved'
    )
    
    if year_filter:
        query = query.filter_by(year=int(year_filter))
    
    if section_filter:
        query = query.filter_by(section=section_filter)
    
    if major_filter:
        query = query.filter_by(major=major_filter)
    
    if search_query:
        query = query.filter(User.student_number.contains(search_query))
    
    members = query.order_by(User.id).all()
    
    filename_parts = [current_user.sig.replace(' ', '_'), 'Member_List']
    
    if year_filter:
        filename_parts.append(f'Year_{year_filter}')
    if section_filter:
        filename_parts.append(f'Section_{section_filter}')
    if major_filter:
        filename_parts.append(f'Major_{major_filter}')
    if search_query:
        filename_parts.append(f'Search_{search_query}')
    
    filename = '_'.join(filename_parts) + f'_{datetime.now().strftime("%Y%m%d")}.docx'

    doc = create_report_document()
    
    title_text = f'{current_user.sig} Member List'
    if any([year_filter, section_filter, major_filter, search_query]):
        title_text += ' (Filtered)'
    
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if any([year_filter, section_filter, major_filter, search_query]):
        filter_info = doc.add_paragraph('Filters Applied:')
        
        if year_filter:
            filter_info.add_run(f' Year: {year_filter}')
        if section_filter:
            filter_info.add_run(f' Section: {section_filter}')
        if major_filter:
            filter_info.add_run(f' Major: {major_filter}')
        if search_query:
            filter_info.add_run(f' Search: {search_query}')
    
    #date
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%Y-%m-%d")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #total mems
    doc.add_paragraph(f'Total Members: {len(members)}')
    doc.add_paragraph()
    
    #create table
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    #headers
    headers = ['No.', 'Student Number', 'Name', 'Year', 'Section', 'Major']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    #member data
    for idx, member in enumerate(members, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = member.student_number or 'N/A'
        row_cells[2].text = member.name
        row_cells[3].text = f'Year {member.year}'
        row_cells[4].text = f'Section {member.section}'
        row_cells[5].text = member.major or 'N/A'
    
    #save
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download/officers_list')
@login_required
def download_officers_list():
    if current_user.user_type != 'admin':
        return redirect(url_for('login'))
    
    #get officers
    officers = Officer.query.filter_by(sig=current_user.sig).order_by(Officer.id).all()
    
    #create docs to download eme eme
    doc = create_report_document()
    
    #title
    title = doc.add_heading(f'{current_user.sig} Officers List', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #date
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%Y-%m-%d")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #total keneme
    doc.add_paragraph(f'Total Officers: {len(officers)}')
    doc.add_paragraph()
    
    #create table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    #headers
    headers = ['No.', 'Position', 'Name']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    #officer data
    for idx, officer in enumerate(officers, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = officer.position
        row_cells[2].text = officer.name
    
    #save
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'{current_user.sig}_Officers_List_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download/events_report')
@login_required
def download_events_report():
    if current_user.user_type != 'admin':
        return redirect(url_for('login'))
    
    #get events
    events = Event.query.filter_by(sig=current_user.sig).order_by(Event.date).all()
    
    #djsalfhnwiebf
    doc = create_report_document()
    
    #eltit
    title = doc.add_heading(f'{current_user.sig} Events Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #da te da te da te
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%Y-%m-%d")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #tutal
    approved_count = len([e for e in events if e.status == 'approved'])
    pending_count = len([e for e in events if e.status == 'pending'])
    rejected_count = len([e for e in events if e.status == 'rejected'])
    
    doc.add_paragraph(f'Total Events: {len(events)}')
    doc.add_paragraph(f'Approved: {approved_count} | Pending: {pending_count} | Rejected: {rejected_count}')
    doc.add_paragraph()
    
    #create table
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    #ULO-ers
    headers = ['No.', 'Title', 'Date', 'Type', 'Status', 'Room']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    #event data
    for idx, event in enumerate(events, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = event.title
        row_cells[2].text = event.date.strftime('%Y-%m-%d')
        row_cells[3].text = event.event_type.title() if event.event_type else 'N/A'
        row_cells[4].text = event.status.title()
        row_cells[5].text = event.room or 'N/A'
    
    #O-save!
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'{current_user.sig}_Events_Report_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download/statistics')
@login_required
def download_statistics():
    if current_user.user_type != 'admin':
        return redirect(url_for('login'))
    
    #get statistics
    members = User.query.filter_by(
        sig=current_user.sig, 
        user_type='student', 
        status='approved'
    ).order_by(User.id).all()
    members_count = len(members)
    officers_count = Officer.query.filter_by(sig=current_user.sig).count()
    events_count = Event.query.filter_by(sig=current_user.sig).count()
    
    #year distrib
    year_counts = {}
    for year in range(1, 5):
        count = User.query.filter_by(sig=current_user.sig, user_type='student', year=year).count()
        year_counts[f'Year {year}'] = count
    
    #paulit ulit nakakasawa HAHAHAHHA
    doc = create_report_document()
    
    # TTT III TTT LLL EEE
    title = doc.add_heading(f'{current_user.sig} Statistics Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # E T A D A T E
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%Y-%m-%d")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # towtal
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f'Total Members: {members_count}')
    doc.add_paragraph(f'Total Officers: {officers_count}')
    doc.add_paragraph(f'Total Events: {events_count}')
    doc.add_paragraph()
    
    #year distrib table
    doc.add_heading('Year Level Distribution', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # heders
    headers = ['Year Level', 'Count', 'Percentage']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    #year day ta
    for year_level, count in year_counts.items():
        percentage = (count / members_count * 100) if members_count > 0 else 0
        row_cells = table.add_row().cells
        row_cells[0].text = year_level
        row_cells[1].text = str(count)
        row_cells[2].text = f'{percentage:.1f}%'
    
    #major distrib
    majors = ['WMAD', 'NETAD', 'AMG', 'SMP', '']
    major_counts = {}
    for major in majors:
        if major == '':
            count = User.query.filter_by(sig=current_user.sig, user_type='student').filter(
                (User.major == '') | (User.major.is_(None))
            ).count()
            major_counts['Not Specified'] = count
        else:
            count = User.query.filter_by(sig=current_user.sig, user_type='student', major=major).count()
            major_counts[major] = count
    
    if members_count > 0:
        doc.add_paragraph()
        doc.add_heading('Major Distribution', level=1)
        table2 = doc.add_table(rows=1, cols=3)
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # wish I was heather~~
        hdr_cells2 = table2.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells2[i].text = header
            hdr_cells2[i].paragraphs[0].runs[0].font.bold = True
        
        #major data
        for major, count in major_counts.items():
            if count > 0:
                percentage = (count / members_count * 100)
                row_cells = table2.add_row().cells
                row_cells[0].text = major
                row_cells[1].text = str(count)
                row_cells[2].text = f'{percentage:.1f}%'
    
    #save
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'{current_user.sig}_Statistics_Report_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download/complete_members')
@login_required
def download_complete_members():
    if current_user.user_type != 'superadmin':
        return redirect(url_for('login'))
    
    #get all mems
    members = User.query.filter_by(user_type='student').order_by(User.sig, User.name).all()
    
    # ito na naman sa pag create ng docs to download
    doc = create_report_document()
    
    # title na naman
    title = doc.add_heading('Complete Member Database', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #      D               A             T               E
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%Y-%m-%d")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #total total total
    doc.add_heading('Summary by SIG', level=1)
    
    sigs = ['CodEx', 'Netac', 'Source Code', 'Robotix', 'Graphicos']
    for sig in sigs:
        count = User.query.filter_by(sig=sig, user_type='student').count()
        doc.add_paragraph(f'{sig}: {count} members')
    
    doc.add_paragraph(f'\nTotal Members: {len(members)}')
    doc.add_paragraph()
    
    #create table
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    #headers - only if you knew
    headers = ['No.', 'SIG', 'Student Number', 'Name', 'Year', 'Section', 'Major']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    #member data - how much I liked you
    for idx, member in enumerate(members, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = member.sig
        row_cells[2].text = member.student_number or 'N/A'
        row_cells[3].text = member.name
        row_cells[4].text = f'Year {member.year}'
        row_cells[5].text = f'Section {member.section}'
        row_cells[6].text = member.major or 'N/A'
    
    #save - but I watched your eyes as he walks by
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'Complete_Member_Database_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download/all_events')
@login_required
def download_all_events():
    if current_user.user_type != 'superadmin':
        return redirect(url_for('login'))
    
    #get all events - what a sight for sore eyes
    events = Event.query.order_by(Event.date).all()
    
    # BRIGHTER THAN THE BLUE SKY
    doc = create_report_document()
    
    #title - he's got you mesmerized while I die
    title = doc.add_heading('All Events Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #date - why would you ever ### me
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%Y-%m-%d")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #total - I'm not even half as pretty
    status_counts = {}
    for event in events:
        status_counts[event.status] = status_counts.get(event.status, 0) + 1
    
    doc.add_paragraph(f'Total Events: {len(events)}')
    for status, count in status_counts.items():
        doc.add_paragraph(f'{status.title()}: {count}')
    doc.add_paragraph()
    
    #create table - you gave him your sweater, it's just polyester
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    #headers - but you like him better
    headers = ['No.', 'SIG', 'Title', 'Date', 'Type', 'Status', 'Room']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    #event data - wish I were _____
    for idx, event in enumerate(events, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = event.sig
        row_cells[2].text = event.title
        row_cells[3].text = event.date.strftime('%Y-%m-%d')
        row_cells[4].text = event.event_type.title() if event.event_type else 'N/A'
        row_cells[5].text = event.status.title()
        row_cells[6].text = event.room or 'N/A'
    
    #save
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'All_Events_Report_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download/sig_stats')
@login_required
def download_sig_stats():
    if current_user.user_type != 'superadmin':
        return redirect(url_for('login'))
    
    #SIG statistics - rome wasn't built in a day
    sigs = ['CodEx', 'Netac', 'Source Code', 'Robotix', 'Graphicos']
    stats = []
    
    for sig in sigs:
        member_count = User.query.filter_by(sig=sig, user_type='student').count()
        officer_count = Officer.query.filter_by(sig=sig).count()
        event_count = Event.query.filter_by(sig=sig).count()
        approved_events = Event.query.filter_by(sig=sig, status='approved').count()
        
        stats.append({
            'sig': sig,
            'members': member_count,
            'officers': officer_count,
            'events': event_count,
            'approved_events': approved_events
        })
    
    total_members = sum(s['members'] for s in stats)
    total_officers = sum(s['officers'] for s in stats)
    total_events = sum(s['events'] for s in stats)
    
    # you gotta climb a little higher, to the top of the display
    doc = create_report_document()
    
    #title - where there's a will, there's a way
    title = doc.add_heading('SIG Comparative Statistics', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    #date - you just gotta slit a throat or two to make the weak obey
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%Y-%m-%d")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # totals - if you want it, just take it
    doc.add_paragraph(f'Total Members Across All SIGs: {total_members}')
    doc.add_paragraph(f'Total Officers Across All SIGs: {total_officers}')
    doc.add_paragraph(f'Total Events Across All SIGs: {total_events}')
    doc.add_paragraph()
    
    #create main table - the world's yours, don't waste it
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    #headers - go make the stars align, to shine BRIGHTER
    headers = ['SIG', 'Members', 'Members %', 'Officers', 'Events', 'Approved Events']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    #SIG data - brighter than the heavens in the skies above
    for stat in stats:
        member_percentage = (stat['members'] / total_members * 100) if total_members > 0 else 0
        row_cells = table.add_row().cells
        row_cells[0].text = stat['sig']
        row_cells[1].text = str(stat['members'])
        row_cells[2].text = f'{member_percentage:.1f}%'
        row_cells[3].text = str(stat['officers'])
        row_cells[4].text = str(stat['events'])
        row_cells[5].text = str(stat['approved_events'])
    
    doc.add_paragraph()
    
    #summary row - you'll be BRIGHTER, goin supernova, all the eyes look up
    summary_row = table.add_row().cells
    summary_row[0].text = 'TOTAL'
    summary_row[0].paragraphs[0].runs[0].font.bold = True
    summary_row[1].text = str(total_members)
    summary_row[2].text = '100%'
    summary_row[3].text = str(total_officers)
    summary_row[4].text = str(total_events)
    summary_row[5].text = str(sum(s['approved_events'] for s in stats))
    
    #save - BRIGHTER
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'SIG_Comparative_Statistics_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/download/system_report')
@login_required
def download_system_report():
    if current_user.user_type != 'superadmin':
        return redirect(url_for('login'))
    
    #get comprehensive system data("COMPREHENSIVE???" HAHAHHAHAH)
    total_students = User.query.filter_by(user_type='student').count()
    total_admins = User.query.filter_by(user_type='admin').count()
    total_events = Event.query.count()
    total_officers = Officer.query.count()
    
    # can't breathe, vision's gettin blurry now
    doc = create_report_document()
    
    # nose bleed, never been so worried
    # how'm I supposed to handle the pressure?
    title = doc.add_heading('SIGMS Complete System Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # it would've been better if I were anybody else
    date_para = doc.add_paragraph(f'{datetime.now().strftime("%Y-%m-%d")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # dear dad, tell me what you would've done
    # too bad, that I'm not your braver son
    doc.add_paragraph('System Administrator: Super Admin')
    doc.add_paragraph()
    
    # why'd I have to be picked to lead?
    # I wish that I could be anybody else
    doc.add_heading('1. Executive Summary', level=1)
    
    doc.add_paragraph(f'Total Students: {total_students}')
    doc.add_paragraph(f'Total Admin Accounts: {total_admins}')
    doc.add_paragraph(f'Total Events: {total_events}')
    doc.add_paragraph(f'Total Officers: {total_officers}')
    doc.add_paragraph(f'Number of SIGs: 5')
    
    doc.add_paragraph()
    
    # I'm supposed to do something
    # am I, am I supposed to do something?
    doc.add_heading('2. SIG Breakdown', level=1)
    
    sigs = ['CodEx', 'Netac', 'Source Code', 'Robotix', 'Graphicos']
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # do I run, do I dare, can I grant my own prayer?
    headers = ['SIG', 'Students', 'Percentage']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # get a grip, grow a pair
    for sig in sigs:
        students = User.query.filter_by(sig=sig, user_type='student').count()
        percentage = (students / total_students * 100) if total_students > 0 else 0
        row_cells = table.add_row().cells
        row_cells[0].text = sig
        row_cells[1].text = str(students)
        row_cells[2].text = f'{percentage:.1f}%'
    
    doc.add_paragraph()
    
    # let me stop you right thereeee
    doc.add_heading('3. Event Status Summary', level=1)
    
    event_statuses = {}
    for event in Event.query.all():
        event_statuses[event.status] = event_statuses.get(event.status, 0) + 1
    
    for status, count in event_statuses.items():
        doc.add_paragraph(f'{status.title()}: {count}')
    
    doc.add_paragraph()
    
    # there's no cause for attack
    doc.add_heading('4. Student Year Distribution', level=1)
    
    year_table = doc.add_table(rows=1, cols=3)
    year_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    year_headers = ['Year Level', 'Count', 'Percentage']
    hdr_cells = year_table.rows[0].cells
    for i, header in enumerate(year_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # swallow your pride, cast your bloodlust aside
    year_dist = {}
    for year in range(1, 5):
        count = User.query.filter_by(user_type='student', year=year).count()
        year_dist[f'Year {year}'] = count
    
    for year_level, count in year_dist.items():
        percentage = (count / total_students * 100) if total_students > 0 else 0
        row_cells = year_table.add_row().cells
        row_cells[0].text = year_level
        row_cells[1].text = str(count)
        row_cells[2].text = f'{percentage:.1f}%'
    
    # bitch, you better fall
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'SIGMS_Complete_System_Report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Logged out successfully!', 'success')
    return redirect(url_for('login'))

def create_default_users():
    """Create default admin and superadmin accounts"""
    #superadmin account
    if not User.query.filter_by(user_type='superadmin').first():
        superadmin = User(
            username='eroldjhonmancilla',
            password=generate_password_hash('mancilla123'),
            user_type='superadmin',
            name='Erold Jhon Mancilla',
            email='eroldjhon.mancilla@lspu.edu.ph'
        )
        db.session.add(superadmin)
    
    #admin accounts
    sigs = ['CodEx', 'Netac', 'Source Code', 'Robotix', 'Graphicos']
    positions = ['director', 'vicedirector']
    
    for sig in sigs:
        for position in positions:
            username = f"{sig.lower().replace(' ', '')}{position}"
            if not User.query.filter_by(username=username).first():
                admin = User(
                    username=username,
                    password=generate_password_hash(f"{sig.lower().replace(' ', '')}{position}"),
                    user_type='admin',
                    name=f"{sig} {position.title()}",
                    email=f"{username}@lspu.edu.ph",
                    sig=sig,
                    position=f"{position.title()} of {sig}"
                )
                db.session.add(admin)
    
    db.session.commit()

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        create_default_users()
    app.run(debug=True)